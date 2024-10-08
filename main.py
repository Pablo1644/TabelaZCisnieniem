# main.py

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import calendar_utils





width = 16.96
height = 0.7
ACTUAL_YEAR = 0
month = 1

def center_table_text(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Tworzenie nowego dokumentu Word
doc = Document()

# Ustawianie marginesów
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(0.75)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Pobieranie danych od użytkownika
month_days = calendar_utils.month_days

while True:
    try:
        ACTUAL_YEAR = int(input("Podaj rok (1951-2050):  "))
        if ACTUAL_YEAR>2050 or ACTUAL_YEAR<1951:
            raise Exception("Rok spoza zakresu. Jeszcze raz podaj dane!")
        month = input("Podaj miesiąc np. Styczeń: ")
        if month not in month_days:
            raise Exception("Nie ma takiego miesiąca! Jeszcze raz podaj dane! ")
        break  # Gdy dane są poprawne, wyjdź z pętli
    except ValueError:
        print("Błędnie podany rok! Spróbuj raz jeszcze.")
    except Exception as e:
        print("Błąd! Podaj dane jescze raz!")  # Wyświetla komunikat, np. "Nie ma takiego miesiąca!"

spec_days = calendar_utils.add_special_days()



if month == 'Luty' and ACTUAL_YEAR % 4 == 0 and (ACTUAL_YEAR % 100 != 0 or ACTUAL_YEAR % 400 == 0):
    month_days['Luty'] = 29

no_of_rows = month_days[month] + 2

# Tworzenie tabeli
table = doc.add_table(no_of_rows, 3)
table.style = 'Table Grid'

# Scalanie pierwszego wiersza
first_row = table.rows[0]
first_row.cells[0].merge(first_row.cells[1]).merge(first_row.cells[2])

# Wymiary komórek
for row in table.rows:
    row.cells[0].width = Cm(width)
    row.cells[1].width = Cm(width)
    row.cells[2].width = Cm(width)

# Wpisanie miesiąca w nagłówku
month_cell = table.rows[0].cells[0]
paragraph = month_cell.paragraphs[0]
run = paragraph.add_run(month.upper() + " " + str(ACTUAL_YEAR))
run.bold = True
run.font.size = Pt(14)

# Nagłówki kolumn
table.rows[1].cells[0].text = "DZIEŃ MIESIĄCA"
table.rows[1].cells[1].text = "DZIEŃ TYGODNIA"
table.rows[1].cells[2].text = "CIŚNIENIE"

# Ustawienie wysokości wierszy
for i in range(2, no_of_rows):
    table.rows[i].height = Cm(height)

# Generowanie dni tygodnia i dni miesiąca
first_day = calendar_utils.first_day_of_the_month(month, ACTUAL_YEAR)
week_days = calendar_utils.get_week_days(first_day, no_of_rows)

# Wypełnianie tabeli
for j in range(2, no_of_rows):
    table.rows[j].cells[0].text = str(j - 1)
    table.rows[j].cells[1].text = week_days[j - 2].upper()
    is_special_day = (j - 1) in spec_days
    calendar_utils.shade_cells_if_needed(table.rows[j], week_days[j - 2], is_special_day)

# Centrowanie tekstu w tabeli
center_table_text(table)

# Zapis dokumentu
doc.save(f'{month}Cisnienie{ACTUAL_YEAR}.docx')
